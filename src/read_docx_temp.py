
import docx
import sys
import os

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Also extract tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading file: {str(e)}"

if __name__ == "__main__":
    file_path = r"d:\Study\研2\开题\开题报告 - 朱军1225.docx"
    content = read_docx(file_path)
    
    # Write to file with utf-8 encoding
    output_path = r"d:\Study\研2\小论文\code\proposal_extracted.txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    
    print(f"Successfully extracted content to {output_path}")
